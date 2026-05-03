"""
多格式导出工具
支持 TXT / Word(.docx) / PDF / EPUB
按PRD要求：可选包含内容标记、保留批注、章节标题格式
"""

import os
import json
import time
from typing import Optional, List


class ExportUtils:
    """多格式导出工具"""

    @staticmethod
    def export_txt(
        chapters: list,
        output_path: str,
        title: str = "未命名小说",
        include_marks: bool = False
    ) -> str:
        """
        导出为TXT文本

        Args:
            chapters: Chapter对象列表(按order排序)
            output_path: 输出文件路径
            title: 小说标题
            include_marks: 是否在文本中包含内容类型标记

        Returns:
            导出文件路径
        """
        lines = []
        lines.append(f"{title}")
        lines.append("=" * 40)
        lines.append(f"导出时间: {time.strftime('%Y-%m-%d %H:%M:%S')}")
        lines.append("")

        for i, ch in enumerate(chapters, 1):
            lines.append(f"第{i}章 {ch.title}")
            lines.append("-" * 30)
            lines.append("")

            if include_marks and ch.content_marks:
                # 在章节末尾追加内容标记说明
                marks_info = []
                for m in ch.content_marks:
                    snippet = ch.content[m.start_pos:min(m.end_pos, m.start_pos + 30)]
                    marks_info.append(f"[{m.content_type}] {snippet}...")
                lines.append(ch.content)
                lines.append("")
                lines.append("【内容标记】")
                for mi in marks_info:
                    lines.append(f"  {mi}")
                lines.append("")
            else:
                lines.append(ch.content)
                lines.append("")

        content = "\n".join(lines)
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)
        return output_path

    @staticmethod
    def export_docx(
        chapters: list,
        output_path: str,
        title: str = "未命名小说",
        include_marks: bool = False,
        keep_annotations: bool = False,
        chapter_title_format: str = "第{num}章 {title}"
    ) -> str:
        """
        导出为Word(.docx)文档

        Args:
            chapters: Chapter对象列表
            output_path: 输出路径
            title: 小说标题
            include_marks: 包含内容标记
            keep_annotations: 保留批注
            chapter_title_format: 章节标题格式模板
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError("请安装 python-docx: pip install python-docx")

        doc = Document()

        # 设置默认字体
        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(12)

        # 总标题
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"导出时间: {time.strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph("")

        for i, ch in enumerate(chapters, 1):
            # 章节标题
            ch_title = chapter_title_format.format(num=i, title=ch.title)
            heading = doc.add_heading(ch_title, level=1)

            # 正文
            para = doc.add_paragraph(ch.content)

            # 内容标记
            if include_marks and ch.content_marks:
                doc.add_paragraph("")
                marks_para = doc.add_paragraph()
                marks_run = marks_para.add_run("【内容标记】")
                marks_run.font.size = Pt(10)
                marks_run.font.color.rgb = RGBColor(128, 128, 128)
                for m in ch.content_marks:
                    snippet = ch.content[m.start_pos:min(m.end_pos, m.start_pos + 50)]
                    m_para = doc.add_paragraph(f"  [{m.content_type}] {snippet}...")
                    m_para.style.font.size = Pt(9)

            # 页分隔
            if i < len(chapters):
                doc.add_page_break()

        doc.save(output_path)
        return output_path

    @staticmethod
    def export_pdf(
        chapters: list,
        output_path: str,
        title: str = "未命名小说",
        include_marks: bool = False,
        chapter_title_format: str = "第{num}章 {title}"
    ) -> str:
        """
        导出为PDF文件

        使用reportlab生成PDF
        """
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import cm
            from reportlab.lib.enums import TA_CENTER, TA_LEFT
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
        except ImportError:
            raise ImportError("请安装 reportlab: pip install reportlab")

        doc = SimpleDocTemplate(
            output_path,
            pagesize=A4,
            rightMargin=2*cm,
            leftMargin=2*cm,
            topMargin=2*cm,
            bottomMargin=2*cm
        )

        styles = getSampleStyleSheet()

        # 自定义样式
        title_style = ParagraphStyle(
            'CustomTitle', parent=styles['Title'],
            fontName='Helvetica-Bold', fontSize=24,
            alignment=TA_CENTER, spaceAfter=30
        )
        chapter_style = ParagraphStyle(
            'ChapterTitle', parent=styles['Heading1'],
            fontName='Helvetica-Bold', fontSize=16,
            spaceBefore=20, spaceAfter=10
        )
        body_style = ParagraphStyle(
            'BodyText2', parent=styles['BodyText'],
            fontName='Helvetica', fontSize=11,
            leading=18, spaceAfter=6
        )
        marks_style = ParagraphStyle(
            'Marks', parent=styles['Normal'],
            fontName='Helvetica', fontSize=8,
            textColor='#888888'
        )

        story = []

        # 总标题
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 0.5*cm))
        story.append(Paragraph(
            f"导出时间: {time.strftime('%Y-%m-%d %H:%M:%S')}",
            body_style
        ))
        story.append(Spacer(1, 1*cm))

        for i, ch in enumerate(chapters, 1):
            ch_title = chapter_title_format.format(num=i, title=ch.title)
            story.append(Paragraph(ch_title, chapter_style))

            # 正文 - 将换行转为<br/>
            content_html = ch.content.replace('\n', '<br/>')
            story.append(Paragraph(content_html, body_style))

            if include_marks and ch.content_marks:
                story.append(Spacer(1, 0.3*cm))
                story.append(Paragraph("【内容标记】", marks_style))
                for m in ch.content_marks:
                    snippet = ch.content[m.start_pos:min(m.end_pos, m.start_pos + 50)]
                    story.append(Paragraph(
                        f"  [{m.content_type}] {snippet}...",
                        marks_style
                    ))

            if i < len(chapters):
                story.append(PageBreak())

        doc.build(story)
        return output_path

    @staticmethod
    def export_epub(
        chapters: list,
        output_path: str,
        title: str = "未命名小说",
        author: str = "未知作者",
        include_marks: bool = False,
        chapter_title_format: str = "第{num}章 {title}"
    ) -> str:
        """
        导出为EPUB电子书
        """
        try:
            from ebooklib import epub
        except ImportError:
            raise ImportError("请安装 ebooklib: pip install ebooklib")

        book = epub.EpubBook()
        book.set_title(title)
        book.set_language('zh-Hans')
        book.add_author(author)

        epub_chapters = []
        spine = ['nav']

        for i, ch in enumerate(chapters, 1):
            ch_title = chapter_title_format.format(num=i, title=ch.title)

            # 构建HTML内容
            content_html = f"<h1>{ch_title}</h1>\n"
            content_html += f"<p>{ch.content.replace(chr(10), '<br/>')}</p>\n"

            if include_marks and ch.content_marks:
                content_html += '<hr/><p style="color:#888;font-size:0.8em;"><strong>内容标记</strong></p>'
                for m in ch.content_marks:
                    snippet = ch.content[m.start_pos:min(m.end_pos, m.start_pos + 50)]
                    content_html += f'<p style="color:#888;font-size:0.7em;">[{m.content_type}] {snippet}...</p>'

            epub_ch = epub.EpubHtml(
                title=ch_title,
                file_name=f'chap_{i:03d}.xhtml',
                lang='zh-Hans'
            )
            epub_ch.content = content_html.encode('utf-8')
            book.add_item(epub_ch)
            epub_chapters.append(epub_ch)
            spine.append(epub_ch)

        # 目录
        book.toc = epub_chapters
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())

        book.spine = spine
        epub.write_epub(output_path, book)
        return output_path

    @staticmethod
    def export(
        format_type: str,
        chapters: list,
        output_path: str,
        **kwargs
    ) -> str:
        """
        统一导出入口

        Args:
            format_type: "TXT" / "DOCX" / "PDF" / "EPUB"
            chapters: 章节列表
            output_path: 输出路径
            **kwargs: 格式特定参数

        Returns:
            导出文件路径
        """
        exporters = {
            "TXT": ExportUtils.export_txt,
            "DOCX": ExportUtils.export_docx,
            "PDF": ExportUtils.export_pdf,
            "EPUB": ExportUtils.export_epub
        }
        exporter = exporters.get(format_type.upper())
        if not exporter:
            raise ValueError(f"不支持的导出格式: {format_type}。支持: {list(exporters.keys())}")
        return exporter(chapters, output_path, **kwargs)
